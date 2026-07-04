# -*- coding: utf-8 -*-
"""
epub_to_docx.py — converte o EPUB Kindle-safe num DOCX que o KDP aceita SEMPRE.

O conversor KFX da Amazon recusa o EPUB mesmo valido (epubcheck 0 erros). O proprio
KDP recomenda DOCX como formato mais confiavel. Le build/<code>/ethics-<code>-kindle.epub
(SVG ja rasterizados) e emite um DOCX preservando o que sobrevive no Kindle reflowable:

  - h1/h2 -> Heading 1, h3 -> Heading 2  (sumario logico do Kindle)
  - SUMARIO embutido na pagina 2 com links internos + marcador "toc" (KDP pede isso)
  - proposicoes (div.prop) -> caixa sombreada + selo em negrito (a voz do Spinoza)
  - demonstracoes (div.demo) -> "Demonstration" em negrito, recuado; QED mantem o ∎
  - glossario-aside (div.gloss) recuado/sombreado; epigrafes em italico centralizado
  - 6 figuras-diagrama embutidas; aberturas/dropcaps/divisores decorativos NAO entram
    (o Kindle reflowa e os descartaria — essa estetica vive no PDF)
  - capa NAO entra no manuscrito (KDP usa a capa enviada separada)

Reusavel pras 6 linguas:  python epub_to_docx.py en  -> build/en/ethics-en.docx
Requisitos: python-docx, lxml.
"""
import io
import re
import sys
import zipfile
from pathlib import Path

from lxml import etree
from lxml import html as lhtml
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass

ROOT = Path(__file__).resolve().parent.parent
BUILD = ROOT / "build"

SKIP_IMG = re.compile(r"(part-opener|divider|/mark\.|tick)")
INLINE_BOLD_CLASS = {"tag", "seal", "num"}      # spans que viram negrito inline
LABEL_CLASSES = {"tag", "seal", "kick", "chapter-num"}
SKIP_DOCS = ("cover.xhtml", "nav.xhtml")
TOC_LABELS = {"en": "Contents", "pt": "Sumário", "es": "Contenido",
              "ru": "Содержание", "zh": "目录", "ja": "目次"}


# ---------- helpers de baixo nivel (XML do Word) ----------
def shade(par, fill="F3EFE4"):
    pPr = par._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def bookmark(par, name, bid):
    s = OxmlElement("w:bookmarkStart")
    s.set(qn("w:id"), str(bid))
    s.set(qn("w:name"), name)
    e = OxmlElement("w:bookmarkEnd")
    e.set(qn("w:id"), str(bid))
    par._p.insert(0, s)
    par._p.append(e)


def link(par, anchor, text):
    h = OxmlElement("w:hyperlink")
    h.set(qn("w:anchor"), anchor)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    st = OxmlElement("w:rStyle")
    st.set(qn("w:val"), "Hyperlink")
    rPr.append(st)
    t = OxmlElement("w:t")
    t.text = text
    r.append(rPr)
    r.append(t)
    h.append(r)
    par._p.append(h)


# ---------- conteudo ----------
def add_inline(par, el, bold=False, italic=False):
    if el.text:
        rn = par.add_run(el.text)
        rn.bold, rn.italic = bold, italic
    for child in el:
        c = set((child.get("class") or "").split())
        b = bold or child.tag in ("strong", "b") or bool(c & INLINE_BOLD_CLASS)
        i = italic or child.tag in ("em", "i")
        if child.tag == "br":
            par.add_run().add_break()
        else:
            add_inline(par, child, b, i)
        if child.tail:
            rn = par.add_run(child.tail)
            rn.bold, rn.italic = bold, italic


def emit_para(doc, el, style=None, align=None, italic=False):
    p = doc.add_paragraph(style=style)
    add_inline(p, el, italic=italic)
    if align is not None:
        p.alignment = align
    return p


def emit_image(doc, z, base, src, caption_el=None):
    path = (base + "/" + src) if base else src
    data = None
    for cand in (path, "OEBPS/" + src, src):
        try:
            data = z.read(cand)
            break
        except KeyError:
            continue
    if data is None:
        return
    doc.add_picture(io.BytesIO(data), width=Inches(4.2))
    doc.paragraphs[-1].alignment = AL.CENTER
    if caption_el is not None:
        emit_para(doc, caption_el, align=AL.CENTER, italic=True)


def cls(el):
    return set((el.get("class") or "").split())


def first(el, xp):
    r = el.xpath(xp)
    return r[0] if r else None


def walk(doc, el, z, base, st):
    """st = {'id': counter, 'titles': [...], 'n': idx} para marcadores do sumario."""
    for node in el:
        tag = node.tag
        c = cls(node)
        if tag in ("h1", "h2", "h3"):
            txt = re.sub(r"^(\d+\.)(\S)", r"\1 \2", ("".join(node.itertext())).strip())
            h = doc.add_heading(txt, level=1 if tag != "h3" else 2)
            if tag in ("h1", "h2"):
                st["n"] += 1
                st["id"] += 1
                bookmark(h, f"toc{st['n']}", st["id"])
        elif tag == "p":
            if "epigraph" in c:
                emit_para(doc, node, align=AL.CENTER, italic=True)
            elif "qed" in c:
                emit_para(doc, node, align=AL.CENTER)
            else:
                emit_para(doc, node)
        elif tag == "blockquote":
            for sub in node:
                if sub.tag == "p":
                    emit_para(doc, sub, style="Quote")
        elif tag == "figure":
            img = first(node, ".//img")
            cap = first(node, ".//figcaption")
            if img is not None:
                emit_image(doc, z, base, img.get("src", ""), cap)
        elif tag == "img":
            src = node.get("src", "")
            if "geo-divider" in c:
                d = doc.add_paragraph()
                d.alignment = AL.CENTER
                d.add_run("❖").font.size = Pt(12)
            elif not SKIP_IMG.search(src):
                emit_image(doc, z, base, src)
        elif tag == "div" and "prop" in c:
            seal = first(node, ".//*[contains(@class,'seal')]")
            if seal is not None:
                sp = doc.add_paragraph()
                sp.alignment = AL.CENTER
                rn = sp.add_run(("".join(seal.itertext())).strip().upper())
                rn.bold = True
                rn.font.size = Pt(9.5)
                rn.font.color.rgb = RGBColor(0x6B, 0x57, 0x2B)
            for tp in node.xpath(".//p"):
                p = emit_para(doc, tp)
                p.paragraph_format.left_indent = Inches(0.3)
                p.paragraph_format.right_indent = Inches(0.3)
                shade(p)
        elif tag == "div" and "demo" in c:
            for tp in node.xpath("./p"):
                p = emit_para(doc, tp)
                p.paragraph_format.left_indent = Inches(0.3)
        elif tag == "div" and "gloss" in c:
            for sub in node:
                if sub.tag == "p":
                    p = emit_para(doc, sub)
                    p.paragraph_format.left_indent = Inches(0.25)
                    shade(p, "EEF1F4")
                elif (sub.tag == "span") and (cls(sub) & LABEL_CLASSES):
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.25)
                    p.add_run(("".join(sub.itertext())).strip()).bold = True
                    shade(p, "EEF1F4")
        elif tag == "dl":
            for sub in node:
                if sub.tag == "dt":
                    p = emit_para(doc, sub)
                    for rn in p.runs:
                        rn.bold = True
                elif sub.tag == "dd":
                    emit_para(doc, sub)
        elif tag in ("ol", "ul"):
            sty = "List Number" if tag == "ol" else "List Bullet"
            for li in node.findall("li"):
                emit_para(doc, li, style=sty)
        elif tag in ("div", "section", "header"):
            if c & LABEL_CLASSES:
                t = ("".join(node.itertext())).strip()
                if t:
                    doc.add_paragraph().add_run(t).bold = True
            else:
                walk(doc, node, z, base, st)


def render_titlepage(doc, body):
    for node in body.iter():
        if node.tag == "h1":
            h = doc.add_heading(("".join(node.itertext())).strip(), level=0)
            h.alignment = AL.CENTER
        elif node.tag == "p":
            p = emit_para(doc, node, align=AL.CENTER)
            if "sub" in cls(node):
                for r in p.runs:
                    r.italic = True


def add_toc_page(doc, code, titles, st):
    hp = doc.add_heading(TOC_LABELS.get(code, "Contents"), level=1)
    st["id"] += 1
    bookmark(hp, "toc", st["id"])
    for i, t in enumerate(titles, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        link(p, f"toc{i}", t)


def base_styles(doc):
    n = doc.styles["Normal"]
    n.font.name = "Georgia"
    n.font.size = Pt(11)
    n.paragraph_format.space_after = Pt(8)
    n.paragraph_format.line_spacing = 1.25


def parse_doc(z, base, href):
    full = (base + "/" + href) if base else href
    root = lhtml.fromstring(z.read(full))
    body = root.find(".//body")
    return body if body is not None else root


def heading_titles(body):
    out = []
    for h in body.xpath(".//h1 | .//h2"):
        out.append(re.sub(r"^(\d+\.)(\S)", r"\1 \2", ("".join(h.itertext())).strip()))
    return out


def convert(code):
    epub = BUILD / code / f"ethics-{code}-kindle.epub"
    if not epub.exists():
        print("nao achei", epub)
        sys.exit(1)
    z = zipfile.ZipFile(epub)
    container = etree.fromstring(z.read("META-INF/container.xml"))
    opf_path = container.find(".//{*}rootfile").get("full-path")
    base = opf_path.rsplit("/", 1)[0] if "/" in opf_path else ""
    opf = etree.fromstring(z.read(opf_path))
    manifest = {it.get("id"): it.get("href") for it in opf.findall(".//{*}item")}
    spine = [manifest[ir.get("idref")] for ir in opf.findall(".//{*}itemref")
             if manifest.get(ir.get("idref")) not in SKIP_DOCS]

    # pre-scan: titulos de nivel 1 (exceto folha de rosto) p/ o sumario
    titles = []
    for href in spine:
        if "titlepage" in href:
            continue
        titles += heading_titles(parse_doc(z, base, href))

    doc = Document()
    base_styles(doc)
    st = {"id": 0, "n": 0}
    started = False
    for href in spine:
        body = parse_doc(z, base, href)
        if "titlepage" in href:
            render_titlepage(doc, body)
            doc.add_page_break()
            add_toc_page(doc, code, titles, st)
            started = True
            continue
        if started:
            doc.add_page_break()
        started = True
        walk(doc, body, z, base, st)

    out = BUILD / code / f"ethics-{code}.docx"
    doc.save(out)
    print("OK DOCX ->", out, f"({out.stat().st_size // 1024} KB) | "
          f"{len(titles)} entradas no sumario")
    return out


if __name__ == "__main__":
    convert(sys.argv[1] if len(sys.argv) > 1 else "en")
